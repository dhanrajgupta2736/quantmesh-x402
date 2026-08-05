import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def markdown_to_docx(md_path, docx_path, title_text="QuantMesh x402 Document"):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x25, 0x2A)

    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    def flush_code():
        nonlocal code_lines
        if code_lines:
            code_text = "".join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(code_text.strip())
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x00, 0x55, 0x88)
            code_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        
        # Parse table markdown
        rows_data = []
        for line in table_lines:
            if re.match(r'^\s*\|?\s*:?-+:?\s*\|', line):
                continue  # Header divider line
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if any(cells):
                rows_data.append(cells)

        if rows_data:
            num_cols = max(len(r) for r in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            for r_idx, row in enumerate(rows_data):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < num_cols:
                        cell = table.cell(r_idx, c_idx)
                        cell.text = cell_text
                        # Style header row
                        if r_idx == 0:
                            set_cell_background(cell, '1E293B')
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        else:
                            bg_color = 'F8FAFC' if r_idx % 2 == 1 else 'FFFFFF'
                            set_cell_background(cell, bg_color)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)
        table_lines = []

    for line in lines:
        raw_line = line.rstrip('\n')

        # Code Block
        if raw_line.strip().startswith('```'):
            if in_code_block:
                flush_code()
                in_code_block = False
            else:
                if in_table:
                    flush_table()
                    in_table = False
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Table
        if raw_line.strip().startswith('|') and '|' in raw_line[1:]:
            in_table = True
            table_lines.append(raw_line)
            continue
        elif in_table:
            flush_table()
            in_table = False

        stripped = raw_line.strip()

        if not stripped:
            continue

        # Headings
        if stripped.startswith('# '):
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(16)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(stripped[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif stripped.startswith('## '):
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(stripped[3:])
            run.font.name = 'Calibri'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
        elif stripped.startswith('### '):
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(3)
            run = h.add_run(stripped[4:])
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        # Blockquote
        elif stripped.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[1:].strip())
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        # Bullet list
        elif stripped.startswith('* ') or stripped.startswith('- ') or stripped.startswith('• '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            content = stripped[2:].strip()
            add_formatted_text(p, content)
        # Numbered list
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            content = re.sub(r'^\d+\.\s', '', stripped).strip()
            add_formatted_text(p, content)
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, stripped)

    if in_code_block:
        flush_code()
    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Successfully created: {docx_path}")

def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    markdown_to_docx(
        'c:/Users/HP/Desktop/quantmesh-x402/docs/ONE_PAGER_HINGLISH.md',
        'c:/Users/HP/Desktop/quantmesh-x402/docs/ONE_PAGER_HINGLISH.docx',
        'QuantMesh x402 1-Page Summary'
    )
    markdown_to_docx(
        'c:/Users/HP/Desktop/quantmesh-x402/docs/PROJECT_EXPLAINER_HINGLISH.md',
        'c:/Users/HP/Desktop/quantmesh-x402/docs/PROJECT_EXPLAINER_HINGLISH.docx',
        'QuantMesh x402 Project Explainer'
    )
    markdown_to_docx(
        'c:/Users/HP/Desktop/quantmesh-x402/PITCH_DECK.md',
        'c:/Users/HP/Desktop/quantmesh-x402/docs/PITCH_DECK.docx',
        'QuantMesh x402 Pitch Deck'
    )
